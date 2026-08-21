"""Live End-to-End integration verification against the running servers."""

import io
import fitz
import httpx
from docx import Document

BASE_URL = "http://127.0.0.1:8000"


def test_live_health():
    print("\n--- 1. Testing GET /health ---")
    res = httpx.get(f"{BASE_URL}/health", timeout=10.0)
    assert res.status_code == 200, f"Expected 200, got {res.status_code}"
    data = res.json()
    print("Health response:", data)
    assert data["status"] == "healthy"
    assert data["tesseract_available"] is True


def test_live_pdf_upload():
    print("\n--- 2. Testing POST /api/analyze with sample PDF ---")
    with open("demo_samples/sample_social_post.pdf", "rb") as f:
        files = {"file": ("sample_social_post.pdf", f, "application/pdf")}
        res = httpx.post(f"{BASE_URL}/api/analyze", files=files, timeout=20.0)
    
    assert res.status_code == 200, f"Failed with {res.status_code}: {res.text}"
    data = res.json()
    print("Extracted filename:", data["file_info"]["filename"])
    print("Extraction method:", data["file_info"]["extraction_method"])
    print("Word count:", data["metrics"]["word_count"])
    print("Engagement Readiness Score:", data["score"]["total_score"], "/ 100")
    print("Verdict:", data["score"]["verdict"])
    print("AI status:", data["ai_review"]["ai_status"])
    print("Strengths count:", len(data["ai_review"]["strengths"]))
    print("Suggestions count:", len(data["ai_review"]["suggestions"]))
    
    assert data["score"]["total_score"] > 0
    assert len(data["ai_review"]["suggestions"]) >= 3
    assert len(data["ai_review"]["improved_post"]) > 20
    return data


def test_live_image_ocr_upload():
    print("\n--- 3. Testing POST /api/analyze with sample PNG (OCR) ---")
    with open("demo_samples/sample_social_post.png", "rb") as f:
        files = {"file": ("sample_social_post.png", f, "image/png")}
        res = httpx.post(f"{BASE_URL}/api/analyze", files=files, timeout=20.0)
    
    assert res.status_code == 200, f"Failed with {res.status_code}: {res.text}"
    data = res.json()
    print("Extracted filename:", data["file_info"]["filename"])
    print("Extraction method:", data["file_info"]["extraction_method"])
    print("Extracted text snippet:", repr(data["extracted_text"][:100]))
    print("Word count:", data["metrics"]["word_count"])
    print("Engagement Readiness Score:", data["score"]["total_score"], "/ 100")
    
    assert "Tesseract OCR" in data["file_info"]["extraction_method"]
    assert data["metrics"]["word_count"] > 10
    assert data["score"]["total_score"] > 0


def test_live_pdf_report_download(payload):
    print("\n--- 4. Testing POST /api/report/pdf ---")
    res = httpx.post(f"{BASE_URL}/api/report/pdf", json=payload, timeout=15.0)
    assert res.status_code == 200, f"PDF report generation failed: {res.text}"
    assert res.headers["content-type"] == "application/pdf"
    assert "attachment; filename=" in res.headers["content-disposition"]
    
    pdf_bytes = res.content
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    assert len(doc) >= 1
    doc_text = doc[0].get_text()
    assert "SOCIAL MEDIA CONTENT ANALYSIS REPORT" in doc_text
    print(f"Generated valid PDF report ({len(pdf_bytes)} bytes, {len(doc)} pages)")
    doc.close()


def test_live_docx_report_download(payload):
    print("\n--- 5. Testing POST /api/report/docx ---")
    res = httpx.post(f"{BASE_URL}/api/report/docx", json=payload, timeout=15.0)
    assert res.status_code == 200, f"DOCX report generation failed: {res.text}"
    assert "wordprocessingml" in res.headers["content-type"]
    assert "attachment; filename=" in res.headers["content-disposition"]
    
    docx_bytes = res.content
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "Social Media Content Analysis Report" in full_text
    print(f"Generated valid DOCX report ({len(docx_bytes)} bytes, {len(doc.paragraphs)} paragraphs)")


if __name__ == "__main__":
    test_live_health()
    payload = test_live_pdf_upload()
    test_live_image_ocr_upload()
    test_live_pdf_report_download(payload)
    test_live_docx_report_download(payload)
    print("\n[SUCCESS] ALL LIVE END-TO-END VERIFICATIONS PASSED SUCCESSFULLY!")
